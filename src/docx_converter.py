"""
src/docx_converter.py — R2 Administrative DOCX Converter per Decree 30/2020/NĐ-CP.

Converts Markdown documents with YAML front matter into styled Word (.docx) files
complying strictly with Vietnamese administrative formatting standards (Nghị định 30/2020/NĐ-CP):
  - Font: Times New Roman throughout
  - Page Margins (A4): Top 20mm, Bottom 20mm, Left 30mm, Right 15mm
  - Header Block:
      * Tên cơ quan ban hành: "ĐẠI HỌC BÁCH KHOA HÀ NỘI" (ALL CAPS, bold/regular)
      * Quốc hiệu: "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" (ALL CAPS, bold, centered)
      * Tiêu ngữ: "Độc lập - Tự do - Hạnh phúc" (Title Case, bold, centered, underlined)
  - Body Text: Line spacing 1.15-1.5, space after 3-6pt, first line indent 1.0-1.27cm, justified
  - Headings & Tables: Clear hierarchy (H1 bold ALL CAPS, H2 bold, H3 italic), tables with shaded headers and grid borders
"""

from __future__ import annotations

import datetime
import os
import re
from pathlib import Path
from typing import Dict, List, Tuple, Any

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor


def parse_front_matter(text: str) -> Tuple[Dict[str, Any], str]:
    """Extract YAML front matter metadata and document body from text.

    Supports flat YAML metadata between `---` delimiters.
    """
    if not text.startswith("---"):
        return {}, text

    lines = text.splitlines()
    closing = None
    for index in range(1, len(lines)):
        if lines[index].strip() == "---":
            closing = index
            break
    if closing is None:
        return {}, text

    fm_block = "\n".join(lines[1:closing])
    body = "\n".join(lines[closing + 1 :]).lstrip("\n")
    return _load_flat_yaml(fm_block), body


def _load_flat_yaml(block: str) -> Dict[str, Any]:
    """Parse flat YAML front matter block into dictionary."""
    try:
        import yaml
        loaded = yaml.safe_load(block) or {}
        if isinstance(loaded, dict):
            return {str(k): v for kk, v in loaded.items() for k in [kk]}
    except Exception:
        pass

    metadata: Dict[str, Any] = {}
    for raw in block.splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or ":" not in line:
            continue
        key, _, value = line.partition(":")
        value = value.split(" #", 1)[0].strip()
        value = value.strip('"').strip("'")
        metadata[key.strip()] = value
    return metadata


def remove_table_borders(table: docx.table.Table) -> None:
    """Remove borders from a Word table for header layout alignment."""
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_grid_borders(table: docx.table.Table, color: str = "000000", sz: str = "4") -> None:
    """Apply thin solid grid borders to a table."""
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_cell_shading(cell: docx.table._Cell, color_hex: str = "EAEAEA") -> None:
    """Set background shading color for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def format_date_str(retrieved_at: str | None = None) -> str:
    """Format date string to Vietnamese administrative date format.
    
    Example: '2026-08-03' -> 'Hà Nội, ngày 03 tháng 08 năm 2026'
    """
    if retrieved_at:
        match = re.search(r"(\d{4})-(\d{1,2})-(\d{1,2})", str(retrieved_at))
        if match:
            year, month, day = match.groups()
            return f"Hà Nội, ngày {int(day):02d} tháng {int(month):02d} năm {year}"

    now = datetime.datetime.now()
    return f"Hà Nội, ngày {now.day:02d} tháng {now.month:02d} năm {now.year}"


def add_formatted_runs(
    paragraph: docx.text.paragraph.Paragraph,
    text: str,
    font_size_pt: float = 13.0,
    default_bold: bool = False,
    default_italic: bool = False,
) -> None:
    """Parse inline markdown tags (**bold**, *italic*, `code`) and add runs to paragraph."""
    # Pattern to match **bold**, *italic*, `code`
    tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        is_bold = default_bold
        is_italic = default_italic
        clean_text = token

        if token.startswith("**") and token.endswith("**") and len(token) >= 4:
            is_bold = True
            clean_text = token[2:-2]
        elif token.startswith("*") and token.endswith("*") and len(token) >= 2:
            is_italic = True
            clean_text = token[1:-1]
        elif token.startswith("`") and token.endswith("`") and len(token) >= 2:
            clean_text = token[1:-1]

        run.bold = is_bold
        run.italic = is_italic
        run.text = clean_text


class Decree30DocxConverter:
    """Converter for creating Decree 30/2020/NĐ-CP compliant Word documents."""

    def __init__(self) -> None:
        pass

    def create_document(self, md_content: str, metadata: Dict[str, Any] | None = None) -> docx.Document:
        parsed_meta, body_text = parse_front_matter(md_content)
        meta = parsed_meta.copy()
        if metadata:
            meta.update(metadata)

        doc = docx.Document()
        self._setup_page_geometry(doc)
        self._setup_default_styles(doc)
        self._add_administrative_header(doc, meta)
        self._add_document_title(doc, meta, body_text)
        self._parse_body(doc, body_text)
        return doc

    def _setup_page_geometry(self, doc: docx.Document) -> None:
        """Set A4 page dimensions and Decree 30 margins: Top 20mm, Bottom 20mm, Left 30mm, Right 15mm."""
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)

    def _setup_default_styles(self, doc: docx.Document) -> None:
        """Set default Normal style parameters throughout the document."""
        normal_style = doc.styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = "Times New Roman"
        normal_font.size = Pt(13)
        normal_font.color.rgb = RGBColor(0, 0, 0)
        
        paragraph_format = normal_style.paragraph_format
        paragraph_format.line_spacing = 1.25
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(4)
        paragraph_format.first_line_indent = Cm(1.27)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_administrative_header(self, doc: docx.Document, meta: Dict[str, Any]) -> None:
        """Add two-column administrative header block per Decree 30/2020/NĐ-CP."""
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Printable width: 16.5cm (21.0 - 3.0 - 1.5)
        # Left col: 7.5cm (~45%), Right col: 9.0cm (~55%)
        table.columns[0].width = Cm(7.5)
        table.columns[1].width = Cm(9.0)
        
        cell_left = table.cell(0, 0)
        cell_right = table.cell(0, 1)

        # Left Column: Agency Name & Document Number
        p1 = cell_left.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.first_line_indent = Cm(0)
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.15
        r1 = p1.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO")
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(12)
        r1.font.color.rgb = RGBColor(0, 0, 0)

        p2 = cell_left.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.first_line_indent = Cm(0)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.15
        r2 = p2.add_run("ĐẠI HỌC BÁCH KHOA HÀ NỘI")
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)
        r2.bold = True
        r2.font.color.rgb = RGBColor(0, 0, 0)

        # Divider line under agency name
        p_div1 = cell_left.add_paragraph()
        p_div1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_div1.paragraph_format.first_line_indent = Cm(0)
        p_div1.paragraph_format.space_before = Pt(0)
        p_div1.paragraph_format.space_after = Pt(4)
        p_div1.paragraph_format.line_spacing = 1.0
        r_div1 = p_div1.add_run("━━━━━━━")
        r_div1.font.name = "Times New Roman"
        r_div1.font.size = Pt(8)
        r_div1.bold = True
        r_div1.font.color.rgb = RGBColor(0, 0, 0)

        p3 = cell_left.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.first_line_indent = Cm(0)
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(0)
        doc_id = meta.get("doc_id", "QĐ-ĐHBK")
        r3 = p3.add_run(f"Số: .../QĐ-ĐHBK")
        r3.font.name = "Times New Roman"
        r3.font.size = Pt(12)
        r3.font.color.rgb = RGBColor(0, 0, 0)

        # Right Column: Motto (Quốc hiệu, Tiêu ngữ) & Date
        p4 = cell_right.paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.first_line_indent = Cm(0)
        p4.paragraph_format.space_before = Pt(0)
        p4.paragraph_format.space_after = Pt(0)
        p4.paragraph_format.line_spacing = 1.15
        r4 = p4.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        r4.font.name = "Times New Roman"
        r4.font.size = Pt(12)
        r4.bold = True
        r4.font.color.rgb = RGBColor(0, 0, 0)

        p5 = cell_right.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.paragraph_format.first_line_indent = Cm(0)
        p5.paragraph_format.space_before = Pt(0)
        p5.paragraph_format.space_after = Pt(2)
        p5.paragraph_format.line_spacing = 1.15
        r5 = p5.add_run("Độc lập - Tự do - Hạnh phúc")
        r5.font.name = "Times New Roman"
        r5.font.size = Pt(13)
        r5.bold = True
        r5.underline = True
        r5.font.color.rgb = RGBColor(0, 0, 0)

        # Divider line under Tiêu ngữ
        p_div2 = cell_right.add_paragraph()
        p_div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_div2.paragraph_format.first_line_indent = Cm(0)
        p_div2.paragraph_format.space_before = Pt(0)
        p_div2.paragraph_format.space_after = Pt(4)
        p_div2.paragraph_format.line_spacing = 1.0
        r_div2 = p_div2.add_run("━━━━━━━━━━━━━━━")
        r_div2.font.name = "Times New Roman"
        r_div2.font.size = Pt(8)
        r_div2.bold = True
        r_div2.font.color.rgb = RGBColor(0, 0, 0)

        p6 = cell_right.add_paragraph()
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p6.paragraph_format.first_line_indent = Cm(0)
        p6.paragraph_format.space_before = Pt(0)
        p6.paragraph_format.space_after = Pt(0)
        date_str = format_date_str(meta.get("retrieved_at"))
        r6 = p6.add_run(date_str)
        r6.font.name = "Times New Roman"
        r6.font.size = Pt(13)
        r6.italic = True
        r6.font.color.rgb = RGBColor(0, 0, 0)

        remove_table_borders(table)

    def _add_document_title(self, doc: docx.Document, meta: Dict[str, Any], body_text: str) -> None:
        """Add main document title below header block."""
        title = meta.get("title", "").strip()
        if not title:
            # Extract first H1 from body if not in metadata
            for line in body_text.splitlines():
                if line.startswith("# "):
                    title = line[2:].strip()
                    break

        if title:
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(6)
            p_space.paragraph_format.space_after = Pt(6)

            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.first_line_indent = Cm(0)
            p_title.paragraph_format.space_before = Pt(12)
            p_title.paragraph_format.space_after = Pt(12)
            p_title.paragraph_format.line_spacing = 1.25

            r_title = p_title.add_run(title.upper())
            r_title.font.name = "Times New Roman"
            r_title.font.size = Pt(15)
            r_title.bold = True
            r_title.font.color.rgb = RGBColor(0, 0, 0)

    def _parse_body(self, doc: docx.Document, body_text: str) -> None:
        """Parse markdown body text into formatted paragraphs, headings, lists, and tables."""
        lines = body_text.splitlines()
        index = 0
        total_lines = len(lines)
        first_h1_skipped = False

        while index < total_lines:
            line = lines[index].rstrip()

            if not line.strip():
                index += 1
                continue

            # Check for Markdown table block
            if line.strip().startswith("|") and "|" in line.strip()[1:]:
                table_lines = []
                while index < total_lines and lines[index].strip().startswith("|"):
                    table_lines.append(lines[index].strip())
                    index += 1
                self._add_table_from_markdown(doc, table_lines)
                continue

            # Headings
            if line.startswith("# "):
                text = line[2:].strip()
                # Skip duplicate main title if already rendered as document title
                if not first_h1_skipped:
                    first_h1_skipped = True
                    index += 1
                    continue
                self._add_heading_h1(doc, text)
                index += 1
                continue

            if line.startswith("## "):
                text = line[3:].strip()
                self._add_heading_h2(doc, text)
                index += 1
                continue

            if line.startswith("### "):
                text = line[4:].strip()
                self._add_heading_h3(doc, text)
                index += 1
                continue

            if line.startswith("#### "):
                text = line[5:].strip()
                self._add_heading_h4(doc, text)
                index += 1
                continue

            # Unordered lists
            match_ul = re.match(r"^(\s*)([-*])\s+(.*)$", line)
            if match_ul:
                indent_str, symbol, text = match_ul.groups()
                level = len(indent_str) // 2
                self._add_list_item(doc, text, is_ordered=False, level=level)
                index += 1
                continue

            # Ordered lists
            match_ol = re.match(r"^(\s*)(\d+[\.\)])\s+(.*)$", line)
            if match_ol:
                indent_str, prefix, text = match_ol.groups()
                level = len(indent_str) // 2
                self._add_list_item(doc, f"{prefix} {text}", is_ordered=True, level=level)
                index += 1
                continue

            # Standard body paragraph
            self._add_body_paragraph(doc, line)
            index += 1

    def _add_heading_h1(self, doc: docx.Document, text: str) -> None:
        """Add H1 Heading (ALL CAPS, Bold, 14pt, Centered)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        add_formatted_runs(p, text.upper(), font_size_pt=14.0, default_bold=True)

    def _add_heading_h2(self, doc: docx.Document, text: str) -> None:
        """Add H2 Heading (Bold, 13pt, Left)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(4)
        add_formatted_runs(p, text, font_size_pt=13.0, default_bold=True)

    def _add_heading_h3(self, doc: docx.Document, text: str) -> None:
        """Add H3 Heading (Italic, 13pt, Left)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        add_formatted_runs(p, text, font_size_pt=13.0, default_italic=True)

    def _add_heading_h4(self, doc: docx.Document, text: str) -> None:
        """Add H4 Heading (Italic, 13pt, Left)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        add_formatted_runs(p, text, font_size_pt=13.0, default_italic=True)

    def _add_list_item(self, doc: docx.Document, text: str, is_ordered: bool = False, level: int = 0) -> None:
        """Add bulleted or numbered list item."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        indent_cm = 0.5 + level * 0.5
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.25

        bullet_prefix = "" if is_ordered else "• "
        add_formatted_runs(p, f"{bullet_prefix}{text}", font_size_pt=13.0)

    def _add_body_paragraph(self, doc: docx.Document, text: str) -> None:
        """Add justified body paragraph with 1.27cm first line indent."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        add_formatted_runs(p, text, font_size_pt=13.0)

    def _add_table_from_markdown(self, doc: docx.Document, table_lines: List[str]) -> None:
        """Convert markdown table lines to a styled docx Table."""
        rows_data = []
        for line in table_lines:
            # Skip separator line (e.g. |---|---|)
            if re.match(r"^\s*\|?\s*[-:]+\s*\|", line):
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows_data.append(cells)

        if not rows_data:
            return

        num_cols = max(len(r) for r in rows_data)
        num_rows = len(rows_data)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        set_table_grid_borders(table, color="000000", sz="4")

        for row_idx, row_content in enumerate(rows_data):
            row_cells = table.rows[row_idx].cells
            is_header = (row_idx == 0)

            for col_idx, cell_value in enumerate(row_content):
                if col_idx < len(row_cells):
                    cell = row_cells[col_idx]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    
                    if is_header:
                        set_cell_shading(cell, "EAEAEA")
                        add_formatted_runs(p, cell_value, font_size_pt=12.0, default_bold=True)
                    else:
                        add_formatted_runs(p, cell_value, font_size_pt=12.0)


def convert_file(md_path: str | Path, output_dir: str | Path = "data/word") -> Path:
    """Convert a single `.md` file into a corresponding Decree 30 `.docx` file."""
    md_path = Path(md_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    target_docx_name = md_path.stem + ".docx"
    target_docx_path = output_dir / target_docx_name

    text = md_path.read_text(encoding="utf-8")
    converter = Decree30DocxConverter()
    doc = converter.create_document(text)
    doc.save(str(target_docx_path))
    return target_docx_path


def convert_all(data_dir: str | Path = "data", output_dir: str | Path = "data/word") -> List[Path]:
    """Convert all harvested `.md` documents in `data_dir` to `.docx` files in `output_dir`."""
    data_path = Path(data_dir)
    output_base = Path(output_dir)
    output_base.mkdir(parents=True, exist_ok=True)

    converted_files: List[Path] = []
    converter = Decree30DocxConverter()

    for path in sorted(data_path.rglob("*.md")):
        if not path.is_file():
            continue

        text = path.read_text(encoding="utf-8")
        doc = converter.create_document(text)

        # Mirror relative subdirectories if any, e.g. data/k3_university/doc.md -> data/word/k3_university/doc.docx
        rel_path = path.relative_to(data_path)
        docx_rel_path = rel_path.with_suffix(".docx")
        
        # Save nested path
        nested_target = output_base / docx_rel_path
        nested_target.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(nested_target))
        converted_files.append(nested_target)

        # If file was in a subfolder, also save a copy directly in data/word/ to ensure flat access
        if len(rel_path.parts) > 1:
            flat_target = output_base / (path.stem + ".docx")
            doc.save(str(flat_target))
            converted_files.append(flat_target)

    return converted_files


if __name__ == "__main__":
    generated = convert_all()
    print(f"Converted {len(generated)} files to Word (.docx) format in data/word/:")
    for f in generated:
        print(f"  - {f} ({f.stat().st_size} bytes)")
