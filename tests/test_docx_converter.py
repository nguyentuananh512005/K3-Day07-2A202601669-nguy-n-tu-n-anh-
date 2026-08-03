"""
tests/test_docx_converter.py — Unit tests for Decree 30 DOCX Converter.
"""

from pathlib import Path
import docx
import pytest

from src.docx_converter import convert_file, convert_all, Decree30DocxConverter, parse_front_matter


def test_parse_front_matter():
    sample = (
        "---\n"
        "doc_id: test-doc\n"
        "title: Test Title Document\n"
        "retrieved_at: 2026-08-03\n"
        "---\n\n"
        "# Section 1\n"
        "This is a test paragraph.\n"
    )
    meta, body = parse_front_matter(sample)
    assert meta.get("doc_id") == "test-doc"
    assert meta.get("title") == "Test Title Document"
    assert body.startswith("# Section 1")


def test_docx_conversion_geometry_and_margins(tmp_path):
    md_file = tmp_path / "sample.md"
    md_file.write_text(
        "---\n"
        "doc_id: sample-id\n"
        "title: Quy định Thử nghiệm\n"
        "retrieved_at: 2026-08-03\n"
        "---\n\n"
        "# Title\n"
        "## 1. Điều khoản thứ nhất\n"
        "Nội dung văn bản thử nghiệm canh đều hai bên.\n",
        encoding="utf-8"
    )

    docx_path = convert_file(md_file, output_dir=tmp_path / "word")
    assert docx_path.exists()
    assert docx_path.stat().st_size > 0

    doc = docx.Document(docx_path)
    section = doc.sections[0]
    
    # Page size A4 (21.0cm x 29.7cm)
    assert abs(section.page_width.cm - 21.0) < 0.1
    assert abs(section.page_height.cm - 29.7) < 0.1

    # Decree 30 Margins: Top 20mm, Bottom 20mm, Left 30mm, Right 15mm
    assert abs(section.top_margin.cm - 2.0) < 0.1
    assert abs(section.bottom_margin.cm - 2.0) < 0.1
    assert abs(section.left_margin.cm - 3.0) < 0.1
    assert abs(section.right_margin.cm - 1.5) < 0.1


def test_administrative_header_content(tmp_path):
    converter = Decree30DocxConverter()
    md_text = (
        "---\n"
        "doc_id: test-header\n"
        "title: Quy định Hành chính HUST\n"
        "retrieved_at: 2026-08-03\n"
        "---\n\n"
        "Nội dung quy định.\n"
    )
    doc = converter.create_document(md_text)
    
    assert len(doc.tables) >= 1
    header_table = doc.tables[0]
    cell_left = header_table.cell(0, 0)
    cell_right = header_table.cell(0, 1)

    left_text = "\n".join(p.text for p in cell_left.paragraphs)
    right_text = "\n".join(p.text for p in cell_right.paragraphs)

    assert "ĐẠI HỌC BÁCH KHOA HÀ NỘI" in left_text
    assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in right_text
    assert "Độc lập - Tự do - Hạnh phúc" in right_text
    assert "Hà Nội, ngày 03 tháng 08 năm 2026" in right_text


def test_convert_all_harvested_documents():
    data_dir = Path("data")
    if not data_dir.exists():
        pytest.skip("data/ directory not found")

    converted = convert_all(data_dir=data_dir, output_dir="data/word")
    assert len(converted) >= 5
    for file in converted:
        assert file.exists()
        assert file.stat().st_size > 0
